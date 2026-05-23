"""
讀 scripts_index.json 的腳本檔，下載內文，跑 match_location() 抽 city+district。
寫到 manual_locations.json（key=case_name, value={city, district, manual:true, source:"script:..."}）。

用法：
  python enrich_from_scripts.py            # 全部 unknown 案件
  python enrich_from_scripts.py weekly     # 只跑本周 weekly_schedule 沒地點的
"""
import io
import json
import re
import sys
from pathlib import Path

from docx import Document
from google.oauth2 import service_account
SA_PATH = "/volume2/docker-prod/scripts/原初映像片庫/service_account.json"
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from supabase import create_client

from extract_locations import match_location, load_env

ROOT = Path(__file__).parent
TOKEN = str(Path(__file__).parent / "token.json")
SCOPES = ["https://www.googleapis.com/auth/drive"]
SCRIPTS_INDEX = ROOT / "scripts_index.json"
MANUAL = ROOT / "manual_locations.json"
UNKNOWN = ROOT / "locations_unknown.json"

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass


def get_service():
    creds = service_account.Credentials.from_service_account_file(SA_PATH, scopes=SCOPES)
    return build("drive", "v3", credentials=creds)


def download_file_bytes(svc, file_id, mime_type):
    """處理 .docx（直接下載）+ Google Doc（export docx）。"""
    if mime_type == "application/vnd.google-apps.document":
        request = svc.files().export_media(
            fileId=file_id,
            mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        request = svc.files().get_media(fileId=file_id)
    buf = io.BytesIO()
    downloader = MediaIoBaseDownload(buf, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    buf.seek(0)
    return buf


def extract_text(file_bytes, mime_type, filename):
    """從 docx / Google Doc / pdf 抽純文字。"""
    name_lower = filename.lower()
    if mime_type == "application/vnd.google-apps.document" or name_lower.endswith(".docx"):
        try:
            doc = Document(file_bytes)
        except Exception as e:
            print(f"  parse fail (docx): {e}", flush=True)
            return ""
        parts = []
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts)
    if name_lower.endswith(".pdf"):
        try:
            import pdfplumber
            file_bytes.seek(0)
            with pdfplumber.open(file_bytes) as pdf:
                return "\n".join((p.extract_text() or "") for p in pdf.pages)
        except Exception as e:
            print(f"  parse fail (pdf): {e}", flush=True)
            return ""
    return ""


def normalize_case(name: str) -> str:
    """跟 sync_schedule.py 對齊，方便比對。"""
    if not name:
        return ""
    name = re.split(r"\s*\n\s*雲端連結", name, maxsplit=1)[0]
    name = re.split(r"\s*\(\s*\d+\s*[-/]", name, maxsplit=1)[0]
    name = re.sub(r"\s+", " ", name).strip()
    return name


def fetch_unknown_cases(client):
    """videos 裡 city is null 的所有 case_name。"""
    PAGE = 1000
    offs = 0
    out = set()
    while True:
        r = client.table("videos").select("case_name").is_("city", "null").range(offs, offs + PAGE - 1).execute()
        rows = r.data or []
        for x in rows:
            cn = (x.get("case_name") or "").strip()
            if cn:
                out.add(cn)
        if len(rows) < PAGE:
            break
        offs += PAGE
    return out


def fetch_weekly_unknown(client):
    """weekly_schedule 裡 city 為空的 case_name。"""
    r = client.table("weekly_schedule").select("case_name,city").execute()
    return {x["case_name"] for x in (r.data or []) if not x.get("city")}


def main():
    only_weekly = "weekly" in sys.argv

    if not SCRIPTS_INDEX.exists():
        print(f"沒有 {SCRIPTS_INDEX}，先跑 scan_scripts.py")
        sys.exit(1)

    env = load_env()
    client = create_client(env["SUPABASE_URL"], env["SUPABASE_SERVICE_ROLE_KEY"])

    if only_weekly:
        target = fetch_weekly_unknown(client)
        print(f"weekly 模式：要查 {len(target)} 個本周案件")
    else:
        target = fetch_unknown_cases(client)
        print(f"全庫模式：要查 {len(target)} 個 city=null 的案件")

    if not target:
        print("沒有要處理的案件")
        return

    scripts = json.loads(SCRIPTS_INDEX.read_text(encoding="utf-8"))["rows"]
    print(f"scripts_index 共 {len(scripts)} 個腳本檔")

    # 索引：case_subfolder（真正案名）→ [scripts]
    by_case = {}
    for s in scripts:
        for key_field in ("case_subfolder", "case_folder"):
            cf = (s.get(key_field) or "").strip()
            if cf:
                by_case.setdefault(cf, []).append(s)

    by_case_norm = {}
    for cf, items in by_case.items():
        by_case_norm.setdefault(normalize_case(cf), []).extend(items)

    # 既存 manual_locations.json
    manual = {}
    if MANUAL.exists():
        manual = json.loads(MANUAL.read_text(encoding="utf-8"))

    svc = get_service()
    hits = 0
    misses = 0
    no_script = 0

    for case_name in sorted(target):
        # 略過已在 manual 裡有的
        if case_name in manual and manual[case_name].get("city"):
            continue

        # 案名本身含明確城市/區名（如「林口空拍」「板橋」）時，直接信案名，
        # 不要被腳本內容覆蓋——腳本可能是早期版本、或同建商其他案的，案名才是當下實拍地點。
        self_loc = match_location(case_name)
        if self_loc and self_loc.get("city") and self_loc.get("district"):
            manual[case_name] = {
                "city": self_loc["city"],
                "district": self_loc["district"],
                "manual": True,
                "source": "case_name",
            }
            hits += 1
            print(f"  ✓ {case_name} → {self_loc['city']} {self_loc['district']} (from case_name)", flush=True)
            continue

        # 找對應的 scripts：先精確比對 case_subfolder，再 normalize 比對
        candidates = by_case.get(case_name, [])
        if not candidates:
            candidates = by_case_norm.get(normalize_case(case_name), [])
        cn_norm = normalize_case(case_name)
        if not candidates and cn_norm:
            # 再試包含關係（資料夾名含 case_name 或反之）
            for cf_norm, items in by_case_norm.items():
                if cn_norm in cf_norm or cf_norm in cn_norm:
                    candidates = items
                    break
        if not candidates and cn_norm:
            # 最後手段：腳本檔名直接含案名
            name_hits = [s for s in scripts if cn_norm in s["name"]]
            if name_hits:
                candidates = name_hits

        if not candidates:
            no_script += 1
            continue

        # 優先挑 .docx / Google Doc
        candidates.sort(
            key=lambda s: (
                0 if s.get("mimeType") == "application/vnd.google-apps.document" else
                1 if s["name"].lower().endswith(".docx") else 9,
                -(s.get("size_bytes") or 0),
            )
        )

        found = None
        for s in candidates[:3]:  # 最多看 3 個腳本
            print(f"  [{case_name}] 試 {s['name']}", flush=True)
            try:
                buf = download_file_bytes(svc, s["drive_file_id"], s.get("mimeType", ""))
                text = extract_text(buf, s.get("mimeType", ""), s["name"])
            except Exception as e:
                print(f"    download/parse fail: {e}", flush=True)
                continue
            if not text.strip():
                continue
            loc = match_location(text)
            if loc and loc.get("city"):
                found = (loc, s["name"])
                break

        if found:
            loc, src = found
            manual[case_name] = {
                "city": loc["city"],
                "district": loc.get("district") or "",
                "manual": True,
                "source": f"script:{src}",
            }
            hits += 1
            print(f"  ✓ {case_name} → {loc['city']} {loc.get('district', '')} (from {src})", flush=True)
        else:
            misses += 1

    MANUAL.write_text(json.dumps(manual, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\n== 完成 ==")
    print(f"找到地點：{hits}")
    print(f"有腳本但內文沒地名：{misses}")
    print(f"完全沒對應腳本：{no_script}")
    print(f"manual_locations.json 已更新（共 {len(manual)} 筆）")


if __name__ == "__main__":
    main()
